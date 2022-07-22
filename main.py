import speech_recognition as sr
import aspose.words as aw
from docx import Document
import time

r = sr.Recognizer()
# create document object
doc = aw.Document()

def listen_voice():
    with sr.Microphone() as source:
        audio = r.listen(source)
        try:
            voice_data = r.recognize_google(audio)
            return voice_data
        except sr.UnknownValueError:
            print('I did not quite catch it')
        except sr.RequestError:
            print('Sorry the servers are down')

def write_file(voice_data):

    # create file name
    ts = str(int(time.time()/17))
    t_file_name = 'text' + ts + '.txt'
    # w_file_name = 'word' + ts + '.docx'
    # wd_file_name = 'word_doc' + ts + '.docx'

    '''
    ISSUE #3
    Cannot overwrite a file using 'w' options
    this leads to entries not saving
    hence using 'a' options
    '''
    # create text file
    fp = open(t_file_name, 'a')

    '''
    ISSUE #1:
    Document is created successfully
    however the entries are in reverse order
    i.e. everything is written bottom to top
    '''
    # # create word file
    # # create a document builder object
    # builder = aw.DocumentBuilder(doc)

    '''
    ISSUE #2:
    Document is created successfully
    however the entries are not saved into the document
    '''
    # # create word using document
    # document = Document()

    # write voice data into file
    if 'document is complete' not in voice_data:
        print(voice_data)

        # wite to text
        fp.write(voice_data + '\n')

        # # write to word
        # builder.writeln(voice_data)

        # # write paragraph using document
        # # document.add_paragraph(voice_data)

        print('written, please continue')
    elif 'document is complete' in voice_data:
        print('closing document')

        '''
        ISSUE #4:
        additional document is created after closing the document
        started happening after introducing timestamp for name
        '''
        # close text document
        fp.close()

        # # save word document
        # doc.save(w_file_name)

        # # save using document
        # document.save(wd_file_name)

        # num = num + 1

        print('create new document or exit?')
        voice_data = listen_voice()
        if 'new document' in voice_data:
            ask()
        else:
            exit()

def ask():
    print('let us create a document... listening')
    while(1):
        voice_data = listen_voice()
        write_file(voice_data)

ask()